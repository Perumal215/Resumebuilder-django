from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.urls import reverse
from .models import Resume
from django.template.loader import render_to_string
from django.http import HttpResponse
from weasyprint import HTML
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import os
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import Resume

def index(request):
    return render(request, 'cvbuilder/index.html')

def about(request):
    return render(request, 'cvbuilder/about.html')

def template_list(request):
    templates = [
        {'id': 1, 'name': 'Modern ', 'image': '/img/template_1.png'},
        {'id': 2, 'name': 'Classic ', 'image': '/img/template_2.png'},
        {'id': 3, 'name': 'Minimal ', 'image': '/img/template3.png'},
        {'id': 4, 'name': 'ElegantEdge ', 'image': '/img/template4.png'},
        {'id': 5, 'name': 'BoldQuote', 'image': '/img/template5.png'},
        {'id': 6, 'name': 'Modern Asymmetrical', 'image': '/img/template6.png'},
        {'id': 7, 'name': 'ModernFocus', 'image': '/img/template7.png'},
        {'id': 8, 'name': 'modernblue', 'image': '/img/tamplate_8.png'},
        {'id': 9, 'name': 'Tech Minimal', 'image': '/img/template9.png'},
        {'id': 10, 'name': 'Corporate Vision', 'image': '/img/templates10.png'},
        {'id': 11, 'name': 'Academic Standard', 'image': '/img/template11.png'},
        {'id': 12, 'name': 'Wilkins Micawber', 'image': '/img/template_12.png'},
        {'id': 13, 'name': 'classicprofessional', 'image': '/img/template-13.png'},
        {'id': 14, 'name': 'ModernSplitColor', 'image': '/img/template-14.png'},
        {'id': 15, 'name': ' Modern Hexagon', 'image': '/img/template15.png'},
        {'id': 16, 'name': 'modern two-column', 'image': '/img/template16.png'},
        {'id': 17, 'name': 'clean_modern', 'image': '/img/template17.png'},
        {'id': 18, 'name': 'Classic Blue Sidebar ', 'image': '/img/template18.png'},
        {'id': 19, 'name': 'Monochrome Modern', 'image': '/img/template19.png'},
        {'id': 20, 'name': 'Vivid Edge Professional', 'image': '/img/template20.png'},
        {'id': 21, 'name': 'Soft Gradient Focus ', 'image': '/img/template21.png'},
        {'id': 22, 'name': 'classic_blue_sidebar', 'image': '/img/template22.png'},
        {'id': 23, 'name': 'modern_serif', 'image': '/img/template23.png'},
        {'id': 24, 'name': 'gradient', 'image': '/img/template24.png'},
        {'id': 25, 'name': 'Emily Jones', 'image': '/img/tamplate_25.png'}]
    
    return render(request, 'cvbuilder/template_list.html', {'templates': templates})

def create_resume(request):
    context = {}
    template_id_raw = request.GET.get('template') or request.POST.get('template_id')
    template_id = int(template_id_raw) if template_id_raw and template_id_raw.isdigit() else None

    if request.method == 'POST':
        image = request.FILES.get('profile_image')
        action = request.POST.get('action')

        if not template_id:
            messages.error(request, "Invalid or missing template. Please select one.")
            return redirect('/')

        # Extract form data (same for all actions)
        personal_info = {
            'full_name': request.POST.get('full_name'),
            'email': request.POST.get('email'),
            'phone': request.POST.get('phone'),
            'address': request.POST.get('address'),
            'linkedin': request.POST.get('linkedin'),
            'website': request.POST.get('website'),
        }

        summary = request.POST.get('summary', '')

        education = [
            {'degree': d, 'field': f, 'institution': i, 'start_date': s, 'end_date': e, 'gpa': g}
            for d, f, i, s, e, g in zip(
                request.POST.getlist('edu_degree[]'),
                request.POST.getlist('edu_field[]'),
                request.POST.getlist('edu_institution[]'),
                request.POST.getlist('edu_start[]'),
                request.POST.getlist('edu_end[]'),
                request.POST.getlist('edu_gpa[]'),
            ) if d or f or i
        ]

        skills = [s.strip() for s in request.POST.get('skills', '').split(',') if s.strip()]

        work_experience = [
            {
                'title': t, 'company': c, 'location': l,
                'start_date': s, 'end_date': e,
                'responsibilities': [r.strip() for r in resp.split('\n') if r.strip()]
            }
            for t, c, l, s, e, resp in zip(
                request.POST.getlist('work_title[]'),
                request.POST.getlist('work_company[]'),
                request.POST.getlist('work_location[]'),
                request.POST.getlist('work_start[]'),
                request.POST.getlist('work_end[]'),
                request.POST.getlist('work_responsibilities[]'),
            ) if t or c or l
        ]

        projects = [
            {'name': n, 'tools': t, 'description': d, 'link': l}
            for n, t, d, l in zip(
                request.POST.getlist('project_name[]'),
                request.POST.getlist('project_tools[]'),
                request.POST.getlist('project_description[]'),
                request.POST.getlist('project_link[]'),
            ) if n or t or d
        ]

        certifications = [
            {'name': n, 'issuer': i, 'date': d}
            for n, i, d in zip(
                request.POST.getlist('cert_name[]'),
                request.POST.getlist('cert_issuer[]'),
                request.POST.getlist('cert_date[]'),
            ) if n or i
        ]

        languages = [
            {'name': n, 'level': l}
            for n, l in zip(
                request.POST.getlist('lang_name[]'),
                request.POST.getlist('lang_level[]'),
            ) if n or l
        ]

        interests = [i.strip() for i in request.POST.get('interests', '').split(',') if i.strip()]

        # ✅ Handle AI Summary Generation (no saving)
        if action == "ai_generate":


            context.update({
                'template_id': template_id,
                'personal_info': personal_info,
                'summary': summary,
                'education': education,
                'skills': ', '.join(skills),
                'work_experience': work_experience,
                'projects': projects,
                'certifications': certifications,
                'languages': languages,
                'interests': ', '.join(interests),
            })

            return render(request, 'cvbuilder/create_resume.html', context)

        # ✅ Only save if action is NOT ai_generate
        resume = Resume.objects.create(
            personal_info=personal_info,
            summary=summary,
            education=education,
            skills=skills,
            work_experience=work_experience,
            projects=projects,
            certifications=certifications,
            awards=[],
            languages=languages,
            publications=[],
            interests=interests,
            references=[],
            profile_image=image,
        )

        if action == 'preview':
            return redirect(reverse('preview_resume', kwargs={'resume_id': resume.id, 'template_id': template_id}))

        # Default: Save only
        messages.success(request, "Resume saved successfully.")
        return render(request, 'cvbuilder/create_resume.html', {
            'template_id': template_id,
            'personal_info': personal_info,
            'summary': summary,
            'education': education,
            'skills': ', '.join(skills),
            'work_experience': work_experience,
            'projects': projects,
            'certifications': certifications,
            'languages': languages,
            'interests': ', '.join(interests),
        })

    # GET request
    return render(request, 'cvbuilder/create_resume.html', {'template_id': template_id})
def preview_resume(request, resume_id, template_id):
    resume = get_object_or_404(Resume, pk=resume_id)
    return render(request, f'template_{template_id}.html', {
        'resume': resume,
        'template_id': template_id,
        'personal_info': resume.personal_info,
        'education': resume.education,
        'skills': resume.skills,
        'summary': resume.summary,
        'work_experience': resume.work_experience,
        'projects': resume.projects,
        'certifications': resume.certifications,
        'awards': resume.awards,
        'languages': resume.languages,
        'publications': resume.publications,
        'interests': resume.interests,
        'references': resume.references,
    })




def edit_resume(request, resume_id, template_id):
    resume = get_object_or_404(Resume, pk=resume_id)

    if request.method == "POST":
        action = request.POST.get("action")

        # Gather form data into Python dicts/lists
        personal_info = {
            "full_name": request.POST.get("full_name", ""),
            "email": request.POST.get("email", ""),
            "phone": request.POST.get("phone", ""),
            "address": request.POST.get("address", ""),
            "linkedin": request.POST.get("linkedin", ""),
            "website": request.POST.get("website", ""),
        }
        summary = request.POST.get("summary", "")

        skills = [s.strip() for s in request.POST.get("skills", "").split(",") if s.strip()]
        interests = [s.strip() for s in request.POST.get("interests", "").split(",") if s.strip()]

        # Build list of dicts for dynamic sections:
        education = []
        for deg, field, inst, start, end, gpa in zip(
            request.POST.getlist("edu_degree[]"),
            request.POST.getlist("edu_field[]"),
            request.POST.getlist("edu_institution[]"),
            request.POST.getlist("edu_start[]"),
            request.POST.getlist("edu_end[]"),
            request.POST.getlist("edu_gpa[]")
        ):
            if deg or field or inst:
                education.append({
                    "degree": deg,
                    "field": field,
                    "institution": inst,
                    "start_date": start,
                    "end_date": end,
                    "gpa": gpa
                })

        work_experience = []
        for title, company, location, start, end, desc in zip(
            request.POST.getlist("work_title[]"),
            request.POST.getlist("work_company[]"),
            request.POST.getlist("work_location[]"),
            request.POST.getlist("work_start[]"),
            request.POST.getlist("work_end[]"),
            request.POST.getlist("work_responsibilities[]")
        ):
            if title or company or location:
                work_experience.append({
                    "title": title,
                    "company": company,
                    "location": location,
                    "start_date": start,
                    "end_date": end,
                    "responsibilities": [line.strip() for line in desc.split("\n") if line.strip()]
                })

        projects = []
        for name, tools, desc, link in zip(
            request.POST.getlist("project_name[]"),
            request.POST.getlist("project_tools[]"),
            request.POST.getlist("project_description[]"),
            request.POST.getlist("project_link[]")
        ):
            if name or tools or desc:
                projects.append({
                    "name": name,
                    "tools": tools,
                    "description": desc,
                    "link": link
                })

        certifications = []
        for name, issuer, date in zip(
            request.POST.getlist("cert_name[]"),
            request.POST.getlist("cert_issuer[]"),
            request.POST.getlist("cert_date[]")
        ):
            if name or issuer:
                certifications.append({
                    "name": name,
                    "issuer": issuer,
                    "date": date
                })

        languages = []
        for lang_name, lang_level in zip(
            request.POST.getlist("lang_name[]"),
            request.POST.getlist("lang_level[]")
        ):
            if lang_name or lang_level:
                languages.append({
                    "name": lang_name,
                    "level": lang_level
                })

        profile_image_file = request.FILES.get("profile_image")

        # If PREVIEW
        if action == "preview":
            context = {
                "template_id": template_id,
                "resume_id": resume_id,
                "personal_info": personal_info,
                "summary": summary,
                "skills": skills,
                "interests": interests,
                "education": education,
                "work_experience": work_experience,
                "projects": projects,
                "certifications": certifications,
                "languages": languages,
                "profile_image_url": resume.profile_image.url if resume.profile_image else None,
            }
            return render(request, f'template_{template_id}.html', context)

        # If SAVE
        elif action == "save":
            resume.personal_info = personal_info
            resume.summary = summary
            resume.skills = skills
            resume.interests = interests
            resume.education = education
            resume.work_experience = work_experience
            resume.projects = projects
            resume.certifications = certifications
            resume.languages = languages
            
            if profile_image_file:
                resume.profile_image = profile_image_file

            resume.save()
            messages.success(request, "Resume updated successfully!")
            return redirect("edit_resume", resume_id=resume.id, template_id=template_id)

    # GET request — pre-fill with DB data
    context = {
        "template_id": template_id,
        "resume_id": resume.id,
        "personal_info": resume.personal_info or {},
        "summary": resume.summary or "",
        "skills": resume.skills or [],
        "interests": resume.interests or [],
        "education": resume.education or [],
        "work_experience": resume.work_experience or [],
        "projects": resume.projects or [],
        "certifications": resume.certifications or [],
        "languages": resume.languages or [],
        "profile_image_url": resume.profile_image.url if resume.profile_image else None,
    }
    return render(request, "cvbuilder/edit_resume.html", context)


def download_resume_pdf(request, resume_id, template_id):
    resume = get_object_or_404(Resume, pk=resume_id)

    html_string = render_to_string(f'template_{template_id}.html', {
        'resume': resume,
        'template_id': template_id, 
        'personal_info': resume.personal_info,
        'education': resume.education,
        'skills': resume.skills,
        'summary': resume.summary,
        'work_experience': resume.work_experience,
        'projects': resume.projects,
        'certifications': resume.certifications,
        'awards': resume.awards,
        'languages': resume.languages,
        'publications': resume.publications,
        'interests': resume.interests,
        'references': resume.references,
        'is_pdf': True,
    })

    pdf_file = HTML(string=html_string, base_url=request.build_absolute_uri()).write_pdf()

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="resume_{resume_id}.pdf"'
    return response

def download_word_resume(request, resume_id, template_id):
    resume = get_object_or_404(Resume, pk=resume_id)
    doc = Document()
    pi = resume.personal_info

    if resume.profile_image and os.path.exists(resume.profile_image.path):
        doc.add_picture(resume.profile_image.path, width=Inches(1.5))

    name_para = doc.add_heading(pi.get('full_name', 'Full Name'), 0)
    name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(f"Email: {pi.get('email', '')}")
    doc.add_paragraph(f"Phone: {pi.get('phone', '')}")
    if pi.get('address'): doc.add_paragraph(f"Address: {pi['address']}")
    if pi.get('linkedin'): doc.add_paragraph(f"LinkedIn: {pi['linkedin']}")
    if pi.get('website'): doc.add_paragraph(f"Website: {pi['website']}")

    if resume.summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(resume.summary)

    if resume.skills:
        doc.add_heading("Skills", level=1)
        for skill in resume.skills:
            doc.add_paragraph(skill, style='List Bullet')

    if resume.education:
        doc.add_heading("Education", level=1)
        for edu in resume.education:
            para = doc.add_paragraph()
            para.add_run(f"{edu.get('degree', '')} in {edu.get('field', '')}").bold = True
            para.add_run(f", {edu.get('institution', '')}")
            doc.add_paragraph(f"{edu.get('start_date', '')} - {edu.get('end_date', '')}")

    if resume.work_experience:
        doc.add_heading("Work Experience", level=1)
        for exp in resume.work_experience:
            para = doc.add_paragraph()
            para.add_run(f"{exp.get('title', '')}").bold = True
            para.add_run(f", {exp.get('company', '')}")
            doc.add_paragraph(f"{exp.get('start_date', '')} - {exp.get('end_date', '')}")
            for line in exp.get('responsibilities', []):
                doc.add_paragraph(f"- {line}", style='List Bullet')

    if resume.projects:
        doc.add_heading("Projects", level=1)
        for proj in resume.projects:
            doc.add_paragraph(f"{proj.get('name', '')} - {proj.get('tools', '')}", style='List Bullet')
            doc.add_paragraph(proj.get('description', ''))

    if resume.certifications:
        doc.add_heading("Certifications", level=1)
        for cert in resume.certifications:
            doc.add_paragraph(f"{cert.get('name', '')} - {cert.get('issuer', '')} ({cert.get('date', '')})", style='List Bullet')

    if resume.languages:
        doc.add_heading("Languages", level=1)
        for lang in resume.languages:
            doc.add_paragraph(f"{lang.get('name', '')} - {lang.get('level', '')}", style='List Bullet')

    if resume.interests:
        doc.add_heading("Interests", level=1)
        doc.add_paragraph(', '.join(resume.interests))

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = pi.get('full_name', 'resume').replace(' ', '_')
    response = HttpResponse(file_stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={filename}_resume.docx'
    return response




